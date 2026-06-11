import os
import json
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, Any, List, Optional

import pandas as pd
import requests
import streamlit as st
from dotenv import load_dotenv

BASE_DIR = Path(__file__).resolve().parent
load_dotenv(BASE_DIR / ".env")

from extractor import (
    DISPLAY_COLUMNS,
    fetch_url,
    extract_pdf_pages,
    extract_html_pages,
    build_output,
    FetchBlockedError,
    FetchParseError,
)
from client_profile_loader import load_client_profile, validate_client_profile
from client_insights_generator import (
    match_research_to_client,
    match_statistics_to_client,
    generate_action_recommendations,
    generate_evidence_backed_action_cards,
    analyze_best_matching_audience_segments,
    generate_content_angles,
    generate_30_60_90_day_plan,
)

APP_TITLE = "Research Extractor"
APP_SUBTITLE = (
    "Upload a PDF, paste a URL, or paste article text. "
    "Review the extracted fields, then send useful results to Teable."
)

CATEGORY_OPTIONS = [
    "Traveler Behavior",
    "Booking Trends",
    "Demographics",
    "Sustainability",
    "Technology & AI",
    "Market Intelligence",
    "Reputation & Reviews",
    "Pricing & Revenue",
    "Content & Social",
    "Seasonality",
    "Wellness",
    "Food & Dining",
    "Luxury",
    "Adventure",
    "Destination competitiveness",
    "Visitor experience",
    "Transportation",
    "Accommodation",
    "Experiences",
    "Food & Beverage",
    "Travel Planning & Booking",
    "Destinations / DMOs",
    "MICE & Business Travel",
    "Special Interest Tourism",
]

DESTINATION_FOCUS_OPTIONS = [
    "Not specified",
    "Global",
    "Europe",
    "Greece",
    "Mediterranean",
    "UK",
    "Germany",
    "USA",
    "Asia Pacific",
    "Middle East",
    "Specific City",
    "Specific Region",
]

TRAVELER_MARKET_OPTIONS = [
    "Business Travelers",
    "Bleisure Travelers",
    "Luxury Travelers",
    "Wellness Travelers",
    "Adventure Travelers",
    "Family Travelers",
    "Solo Travelers",
    "Group Travelers",
    "Leisure Travelers",
]

RESEARCH_TYPE_OPTIONS = [
    "Report",
    "Survey",
    "Article",
    "eBook",
    "Infographic",
    "Whitepaper",
    "Case Study",
]

STATUS_OPTIONS = [
    "New",
    "Needs Review",
    "Approved",
    "Sent to Content",
    "Used in Strategy",
    "Archived",
    "Rejected",
]

SOURCE_QUALITY_OPTIONS = [
    "High",
    "Medium",
    "Low",
    "Unclear",
]

USEFUL_FOR_OPTIONS = [
    "SEO",
    "Social Media",
    "Paid Ads",
    "Content Strategy",
    "Branding",
    "PR",
    "Email Marketing",
    "Market Research",
    "Client Proposal",
    "Strategy Deck",
    "Website Content",
    "Campaign Planning",
]

RELEVANT_CLIENT_TYPES_OPTIONS = [
    "Hotels",
    "Luxury Hotels",
    "Villas",
    "DMOs",
    "Tour Operators",
    "Travel Agencies",
    "Restaurants",
    "Cruises",
    "Airlines",
    "Car Rentals",
    "Experiences Providers",
]

TREND_STRENGTH_OPTIONS = [
    "Weak Signal",
    "Growing Trend",
    "Strong Trend",
    "Market Shift",
]

TEABLE_FIELD_MAP = {
    "Title": "Title",
    "Publisher": "Publisher",
    "Date": "Date",
    "Sample": "Sample",
    "Methodology": "Methodology",
    "Research Type": "RESEARCH_TYPE",
    "Category": "CATEGORY",
    "Destination Focus": "DESTINATION_FOCUS",
    "Traveler Market": "TRAVELER_MARKET",
    "Data Points": "Data Points",
    "Summary": "Summary",
    "Conclusion": "Conclusion",
    "Digital Marketing Insight": "Digital Marketing Insight",
    "Usefulness": "Usefulness",
    "Client-ready Insight": "Client-ready Insight",
    "Content Ideas": "Content Ideas",
    "Key Statistics": "Key Statistics",
    "Useful For": "Useful For",
    "Relevant Client Types": "Relevant Client Types",
    "Trend Strength": "Trend Strength",
    "Research Value Score": "Research Value Score",
}

OPTIONAL_TEABLE_FIELD_MAP = {
    "Ethnicity Focus": "Ethnicity Focus",
    "PDF Source URL": "PDF Source URL",
    "PDF File Name": "PDF File Name",
    "Source Type": "Source Type",
    "OCR Used": "OCR Used",
    "Gemini Used": "Gemini Used",
    "Confidence Score": "Confidence Score",
    "Needs Review Count": "Needs Review Count",
    "Status": "Status",
    "Source Quality": "Source Quality",
    "Reviewer Notes": "Reviewer Notes",
}


def get_secret(name: str, default: str = "") -> str:
    env_value = os.getenv(name, "").strip()
    if env_value:
        return env_value

    try:
        value = st.secrets.get(name, default)
        return str(value).strip() if value is not None else default
    except Exception:
        return default


def require_password() -> bool:
    password = get_secret("APP_PASSWORD")
    if not password:
        return True

    if st.session_state.get("app_authenticated"):
        return True

    st.title(APP_TITLE)
    st.caption("Private research extraction tool")
    entered = st.text_input("Password", type="password")
    if st.button("Enter", use_container_width=True):
        if entered == password:
            st.session_state.app_authenticated = True
            st.rerun()
        else:
            st.error("Incorrect password.")

    return False


def _clean_teable_value(value: Any) -> Any:
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return value
    text = str(value or "").strip()
    if not text or text.lower() == "not specified":
        return ""
    return text


def _teable_multi_select_value(value: Any) -> List[str]:
    """Convert app semicolon/comma/newline separated multi-values to a Teable multiple-select array."""
    if value in ("", None):
        return []
    if isinstance(value, list):
        return [str(v).strip() for v in value if str(v).strip() and str(v).strip().lower() != "not specified"]

    text = str(value or "").strip()
    if not text or text.lower() == "not specified":
        return []

    raw_parts = []
    for chunk in text.replace("\n", ";").split(";"):
        part = chunk.strip()
        if part:
            raw_parts.append(part)

    if len(raw_parts) <= 1 and "," in text:
        raw_parts = [p.strip() for p in text.split(",") if p.strip()]

    deduped = []
    seen = set()
    for item in raw_parts:
        if item.lower() == "not specified":
            continue
        key = item.lower()
        if key not in seen:
            seen.add(key)
            deduped.append(item)
    return deduped

def build_teable_fields(final_row: Dict[str, Any], source_url: str = "") -> Dict[str, Any]:
    # Force-send extractor-generated insight fields to Teable.
    # If Usefulness is empty, derive it from the generated insight fields.
    if not str(final_row.get("Digital Marketing Insight", "")).strip():
        final_row["Digital Marketing Insight"] = str(final_row.get("Client-ready Insight", "") or "").strip()

    if not str(final_row.get("Usefulness", "")).strip():
        usefulness_parts = []
        for source_field in ["Client-ready Insight", "Digital Marketing Insight", "Summary", "Key Statistics"]:
            value = str(final_row.get(source_field, "") or "").strip()
            if value and value.lower() != "not specified":
                usefulness_parts.append(value)
        final_row["Usefulness"] = "\n\n".join(usefulness_parts[:3])

    fields: Dict[str, Any] = {}

    multiple_select_fields = {"Category", "Traveler Market", "Useful For", "Relevant Client Types"}

    for teable_field, app_field in TEABLE_FIELD_MAP.items():
        value = final_row.get(app_field, "")
        if teable_field in multiple_select_fields:
            fields[teable_field] = _teable_multi_select_value(value)
        else:
            fields[teable_field] = _clean_teable_value(value)

    for teable_field, app_field in OPTIONAL_TEABLE_FIELD_MAP.items():
        value = final_row.get(app_field, "")
        if value not in ("", None):
            if teable_field in multiple_select_fields:
                fields[teable_field] = _teable_multi_select_value(value)
            else:
                fields[teable_field] = _clean_teable_value(value)

    fields["Source URL"] = source_url.strip()
    fields["Created At"] = datetime.now(timezone.utc).isoformat()
    return fields


def _teable_headers() -> Dict[str, str]:
    api_token = get_secret("TEABLE_API_TOKEN")
    if not api_token:
        raise RuntimeError("TEABLE_API_TOKEN is missing")
    return {
        "Authorization": f"Bearer {api_token}",
        "Content-Type": "application/json",
    }


def _teable_endpoint(path: str) -> str:
    api_url = get_secret("TEABLE_API_URL", "https://app.teable.ai").rstrip("/")
    return f"{api_url}{path}"



def list_teable_records(limit: int = 100) -> List[Dict[str, Any]]:
    table_id = get_secret("TEABLE_TABLE_ID")
    if not table_id:
        return []

    endpoint = _teable_endpoint(f"/api/table/{table_id}/record")
    params = {
        "fieldKeyType": "name",
        "take": str(limit),
    }

    response = requests.get(endpoint, headers=_teable_headers(), params=params, timeout=30)
    if response.status_code >= 400:
        return []

    data = response.json()
    records = data.get("records") if isinstance(data, dict) else None
    return records if isinstance(records, list) else []


def find_duplicate_in_teable(final_row: Dict[str, Any], source_url: str = "") -> Optional[Dict[str, Any]]:
    title = str(final_row.get("Title", "")).strip().lower()
    url = str(source_url or "").strip().lower()

    if not title and not url:
        return None

    for record in list_teable_records(limit=150):
        fields = record.get("fields", {}) if isinstance(record, dict) else {}
        existing_title = str(fields.get("Title", "")).strip().lower()
        existing_url = str(fields.get("Source URL", "")).strip().lower()

        if url and existing_url and url == existing_url:
            return record
        if title and existing_title and title == existing_title:
            return record

    return None


def append_result_to_teable(final_row: Dict[str, Any], source_url: str = "") -> Dict[str, Any]:
    table_id = get_secret("TEABLE_TABLE_ID")
    if not table_id:
        raise RuntimeError("TEABLE_TABLE_ID is missing")

    endpoint = _teable_endpoint(f"/api/table/{table_id}/record")
    fields = build_teable_fields(final_row, source_url)
    payload = {
        "fieldKeyType": "name",
        "typecast": True,
        "records": [{"fields": fields}],
    }

    response = requests.post(endpoint, json=payload, headers=_teable_headers(), timeout=30)

    # If optional metadata fields have not been created in Teable yet,
    # retry once with only the core fields.
    if response.status_code >= 400 and any(k in fields for k in OPTIONAL_TEABLE_FIELD_MAP):
        core_fields = {
            k: v
            for k, v in fields.items()
            if k not in OPTIONAL_TEABLE_FIELD_MAP and k not in {"Ethnicity Focus", "Digital Marketing Insight", "PDF Source URL", "PDF File Name", "Source Type", "OCR Used", "Gemini Used", "Confidence Score", "Needs Review Count", "Status", "Source Quality", "Usefulness", "Reviewer Notes"}
        }
        payload["records"][0]["fields"] = core_fields
        response = requests.post(endpoint, json=payload, headers=_teable_headers(), timeout=30)

    if response.status_code >= 400:
        try:
            detail = response.json()
        except Exception:
            detail = response.text
        raise RuntimeError(f"Teable API error {response.status_code}: {detail}")

    try:
        return response.json()
    except Exception:
        return {"status": "ok"}


def friendly_llm_error(message: str) -> str:
    msg = (message or "").lower()

    if (
        "quota was exhausted" in msg
        or "quota exhausted" in msg
        or "resource_exhausted" in msg
        or "429" in msg
        or "daily limit" in msg
    ):
        return "You hit the daily limit."

    if "503" in msg or "unavailable" in msg or "high demand" in msg:
        return "Gemini is temporarily unavailable. Please try again later."

    return message or "Gemini error."


def split_multi(value: str) -> List[str]:
    if not value or str(value).strip().lower() == "not specified":
        return []
    parts = [p.strip() for p in str(value).replace("\n", ";").split(";")]
    return [p for p in parts if p]


def join_multi(values: List[str]) -> str:
    return "; ".join(values) if values else "Not specified"


def normalize_choice(value: str, options: List[str], default: str = "Not specified") -> str:
    value = str(value or "").strip()
    return value if value in options else default


def review_to_dataframe(final_row: Dict[str, Any]) -> pd.DataFrame:
    return pd.DataFrame(
        [{"Field": field, "Value": final_row.get(field, "Not specified")} for field in DISPLAY_COLUMNS]
    )


def row_to_dataframe(final_row: Dict[str, Any]) -> pd.DataFrame:
    ordered = {field: final_row.get(field, "Not specified") for field in DISPLAY_COLUMNS}
    return pd.DataFrame([ordered])


def review_meta_to_dataframe(review_meta: Dict[str, Dict[str, Any]]) -> pd.DataFrame:
    rows: List[Dict[str, Any]] = []
    for field in DISPLAY_COLUMNS:
        meta = review_meta.get(field, {})
        rows.append(
            {
                "Field": field,
                "Value": meta.get("value", ""),
                "Confidence %": meta.get("confidence_pct", 0),
                "Needs review": meta.get("needs_review", True),
                "Method": meta.get("extraction_method", ""),
                "Evidence page": meta.get("evidence_page", ""),
                "Evidence snippet": meta.get("evidence_snippet", ""),
            }
        )
    return pd.DataFrame(rows)


def style_not_specified(df: pd.DataFrame):
    def highlight(val):
        if isinstance(val, str) and val.strip().lower() == "not specified":
            return "color: #ffb3b3;"
        return ""

    return df.style.map(highlight, subset=["Value"])


def extract_pasted_article_pages(article_text: str) -> List[Dict[str, Any]]:
    cleaned = article_text.strip()
    if not cleaned:
        raise ValueError("Pasted article text is empty.")
    return extract_html_pages(cleaned)


def compute_source_quality(row: Dict[str, Any], avg_conf: int) -> str:
    publisher_ok = str(row.get("Publisher", "")).strip().lower() not in {"", "not specified"}
    methodology_ok = str(row.get("Methodology", "")).strip().lower() not in {"", "not specified"}
    data_ok = str(row.get("Data Points", "")).strip().lower() not in {"", "not specified"}
    conclusion_ok = str(row.get("Conclusion", "")).strip().lower() not in {"", "not specified"}

    if avg_conf >= 85 and publisher_ok and methodology_ok and data_ok:
        return "High"
    if avg_conf >= 75 and publisher_ok and (data_ok or conclusion_ok):
        return "Medium"
    if avg_conf < 65:
        return "Low"
    return "Unclear"


def compute_default_status(needs_review_count: int) -> str:
    return "Needs Review" if needs_review_count > 0 else "New"


def generate_default_usefulness(row: Dict[str, Any], metadata: Dict[str, Any]) -> str:
    title = str(row.get("Title", "")).strip()
    category = str(row.get("CATEGORY", "")).strip()
    traveler_market = str(row.get("TRAVELER_MARKET", "")).strip()
    digital_insight = str(row.get("Digital Marketing Insight", "")).strip()
    summary = str(row.get("Summary", "")).strip()
    data_points = str(row.get("Data Points", "")).strip()

    parts = []

    if title and title.lower() != "not specified":
        parts.append(f"Useful source for tracking insights from '{title}'.")

    if category and category.lower() != "not specified":
        parts.append(f"It can support monitoring of themes such as {category}.")

    if traveler_market and traveler_market.lower() != "not specified":
        parts.append(f"It is relevant for understanding traveler segments such as {traveler_market}.")

    if digital_insight and digital_insight.lower() != "not specified":
        parts.append(f"Marketing relevance: {digital_insight}")

    elif summary and summary.lower() != "not specified":
        parts.append(f"Strategic relevance: {summary}")

    if data_points and data_points.lower() != "not specified":
        first_dp = data_points.replace("•", "").splitlines()[0].strip()
        if first_dp:
            parts.append(f"Contains usable supporting data, for example: {first_dp}")

    return " ".join(parts).strip()

def render_editable_review(row: Dict[str, Any], metadata: Dict[str, Any]) -> Dict[str, Any]:
    st.subheader("Review before sending")
    st.caption("Edit any field before sending the result to Teable.")

    edited: Dict[str, Any] = {}

    col1, col2, col3 = st.columns(3)
    with col1:
        edited["Title"] = st.text_input("Title", value=str(row.get("Title", "")))
    with col2:
        edited["Publisher"] = st.text_input("Publisher", value=str(row.get("Publisher", "")))
    with col3:
        edited["Date"] = st.text_input("Date", value=str(row.get("Date", "")))

    default_categories = [x for x in split_multi(row.get("CATEGORY", "")) if x in CATEGORY_OPTIONS]
    default_travelers = [x for x in split_multi(row.get("TRAVELER_MARKET", "")) if x in TRAVELER_MARKET_OPTIONS]

    edited["CATEGORY"] = join_multi(
        st.multiselect("CATEGORY", CATEGORY_OPTIONS, default=default_categories)
    )

    d_col1, d_col2 = st.columns(2)
    with d_col1:
        dest_default = normalize_choice(row.get("DESTINATION_FOCUS", ""), DESTINATION_FOCUS_OPTIONS, "Not specified")
        edited["DESTINATION_FOCUS"] = st.selectbox(
            "DESTINATION_FOCUS",
            DESTINATION_FOCUS_OPTIONS,
            index=DESTINATION_FOCUS_OPTIONS.index(dest_default),
        )
    with d_col2:
        type_default = normalize_choice(row.get("RESEARCH_TYPE", ""), RESEARCH_TYPE_OPTIONS, "Report")
        edited["RESEARCH_TYPE"] = st.selectbox(
            "RESEARCH_TYPE",
            RESEARCH_TYPE_OPTIONS,
            index=RESEARCH_TYPE_OPTIONS.index(type_default),
        )

    edited["TRAVELER_MARKET"] = join_multi(
        st.multiselect("TRAVELER_MARKET", TRAVELER_MARKET_OPTIONS, default=default_travelers)
    )

    edited["Ethnicity Focus"] = st.text_input(
        "Ethnicity Focus",
        value=str(row.get("Ethnicity Focus", "")),
        help="Ethnicity, nationality, demographic, or source-market focus if the source explicitly mentions one.",
    )

    edited["Sample"] = st.text_area("Sample", value=str(row.get("Sample", "")), height=90)
    edited["Methodology"] = st.text_area("Methodology", value=str(row.get("Methodology", "")), height=120)
    edited["Data Points"] = st.text_area("Data Points", value=str(row.get("Data Points", "")), height=170)
    edited["Summary"] = st.text_area(
        "Summary",
        value=str(row.get("Summary", "")),
        height=130,
        help="General 2-4 sentence summary of the source.",
    )
    edited["Conclusion"] = st.text_area(
        "Conclusion",
        value=str(row.get("Conclusion", "")),
        height=110,
        help="Final takeaway, implication, or recommendation. Separate from Summary.",
    )
    edited["Digital Marketing Insight"] = st.text_area(
        "Digital Marketing Insight",
        value=str(row.get("Digital Marketing Insight", "")),
        height=120,
        placeholder="Useful input, advice, campaign idea, content/social/SEO hint, or practical marketing hack from the source.",
    )

    st.subheader("Client & research intelligence")
    edited["Client-ready Insight"] = st.text_area(
        "Client-ready Insight",
        value=str(row.get("Client-ready Insight", "")),
        height=110,
        placeholder="Concise insight that can be sent to a client or used in a client-facing report.",
    )
    edited["Content Ideas"] = st.text_area(
        "Content Ideas",
        value=str(row.get("Content Ideas", "")),
        height=130,
        placeholder="Content, campaign, SEO, PR, social or strategy ideas.",
    )
    edited["Key Statistics"] = st.text_area(
        "Key Statistics",
        value=str(row.get("Key Statistics", "")),
        height=130,
        placeholder="Key statistics, percentages, scores or numeric findings.",
    )

    useful_default = [x for x in split_multi(row.get("Useful For", "")) if x in USEFUL_FOR_OPTIONS]
    client_types_default = [x for x in split_multi(row.get("Relevant Client Types", "")) if x in RELEVANT_CLIENT_TYPES_OPTIONS]
    edited["Useful For"] = join_multi(st.multiselect("Useful For", USEFUL_FOR_OPTIONS, default=useful_default))
    edited["Relevant Client Types"] = join_multi(st.multiselect("Relevant Client Types", RELEVANT_CLIENT_TYPES_OPTIONS, default=client_types_default))

    t_col1, t_col2 = st.columns(2)
    with t_col1:
        trend_default = normalize_choice(row.get("Trend Strength", ""), TREND_STRENGTH_OPTIONS, "Growing Trend")
        edited["Trend Strength"] = st.selectbox(
            "Trend Strength",
            TREND_STRENGTH_OPTIONS,
            index=TREND_STRENGTH_OPTIONS.index(trend_default),
        )
    with t_col2:
        try:
            score_default = int(float(row.get("Research Value Score", 0)))
        except Exception:
            score_default = 0
        edited["Research Value Score"] = st.number_input(
            "Research Value Score",
            min_value=0,
            max_value=100,
            value=max(0, min(100, score_default)),
            step=1,
        )

    st.subheader("Workflow metadata")
    m1, m2, m3 = st.columns(3)
    with m1:
        edited["Status"] = st.selectbox(
            "Status",
            STATUS_OPTIONS,
            index=STATUS_OPTIONS.index(metadata.get("Status", "Needs Review")),
        )
    with m2:
        edited["Source Quality"] = st.selectbox(
            "Source Quality",
            SOURCE_QUALITY_OPTIONS,
            index=SOURCE_QUALITY_OPTIONS.index(metadata.get("Source Quality", "Unclear")),
        )
    with m3:
        edited["Source Type"] = st.text_input("Source Type", value=str(metadata.get("Source Type", "")))

    edited["PDF Source URL"] = st.text_input(
        "PDF Source URL",
        value=str(metadata.get("PDF Source URL", "")),
        help="Clickable source URL for the PDF/article. In Teable, create this as a URL field.",
    )
    edited["PDF File Name"] = st.text_input(
        "PDF File Name",
        value=str(metadata.get("PDF File Name", "")),
        help="Uploaded local PDF filename. This is not clickable unless the file also has a public URL.",
    )

    edited["Usefulness"] = st.text_area(
        "Usefulness",
        value=str(metadata.get("Usefulness", "")),
        height=90,
        placeholder="Why is this source useful for content, marketing, strategy, or trend monitoring?",
    )
    edited["Reviewer Notes"] = st.text_area(
        "Reviewer Notes",
        value=str(metadata.get("Reviewer Notes", "")),
        height=90,
        placeholder="Manual notes, issues, or follow-up tasks.",
    )

    edited["OCR Used"] = bool(metadata.get("OCR Used", False))
    edited["Gemini Used"] = bool(metadata.get("Gemini Used", False))
    edited["Confidence Score"] = int(metadata.get("Confidence Score", 0))
    edited["Needs Review Count"] = int(metadata.get("Needs Review Count", 0))

    return edited



def load_research_blocks(path: str = "data/research_blocks.json") -> List[Dict[str, Any]]:
    """Load reusable research blocks for client insight matching."""
    research_path = BASE_DIR / path

    with open(research_path, "r", encoding="utf-8") as file:
        return json.load(file)



def build_audience_intelligence_markdown(
    client_profile: Dict[str, Any],
    matches: List[Dict[str, Any]],
    matched_statistics: List[Dict[str, Any]],
    action_recommendations: Dict[str, List[str]] = None,
    action_cards: List[Dict[str, Any]] = None,
    audience_segments: List[Dict[str, Any]] = None,
    content_angles: List[Dict[str, Any]] = None,
    day_plan: Dict[str, List[str]] = None,
) -> str:
    """
    Build a simple Markdown export focused only on traveler data points
    and the research sources behind them. No actions, content pillars,
    digital marketing plan or 30/60/90 plan are included.
    """
    client_name = str(client_profile.get("client_name", "Client") or "Client").strip()
    website = str(client_profile.get("website", "") or "").strip()
    main_destination = str(client_profile.get("main_destination", "") or "").strip()
    main_vertical = str(client_profile.get("main_vertical", "") or "").strip()

    def join_values(values: Any) -> str:
        if isinstance(values, list):
            cleaned = [str(value).strip() for value in values if str(value).strip()]
            return ", ".join(cleaned) if cleaned else "Not specified"
        text_value = str(values or "").strip()
        return text_value if text_value else "Not specified"

    def source_label(source: Dict[str, Any]) -> str:
        source = source or {}
        name = str(source.get("source_name", "") or "—").strip()
        report = str(source.get("source_report", "") or "").strip()
        year = str(source.get("source_year", "") or "").strip()
        if report and year:
            return f"{name} — {report} ({year})"
        if report:
            return f"{name} — {report}"
        if year:
            return f"{name} ({year})"
        return name

    lines: List[str] = []
    lines.append(f"# {client_name} — Traveler Data Points")
    lines.append("")
    lines.append("## Client / Partner Profile")
    lines.append("")
    lines.append(f"- **Website:** {website or 'Not specified'}")
    lines.append(f"- **Main destination:** {main_destination or 'Not specified'}")
    lines.append(f"- **Main vertical:** {main_vertical or 'Not specified'}")
    lines.append(f"- **Products / services:** {join_values(client_profile.get('products_services', []))}")
    lines.append(f"- **Target audiences:** {join_values(client_profile.get('target_audiences', []))}")
    lines.append(f"- **Markets:** {join_values(client_profile.get('markets', []))}")
    lines.append(f"- **Research needs:** {join_values(client_profile.get('research_needs', []))}")
    lines.append("")
    lines.append("---")
    lines.append("")

    lines.append("## 1. Matched Traveler Data Points")
    lines.append("")
    lines.append("Τα παρακάτω στοιχεία επιλέχθηκαν επειδή ταιριάζουν με το προφίλ, το προϊόν, το κοινό ή τις αγορές ενδιαφέροντος.")
    lines.append("")

    if not matched_statistics:
        lines.append("_Δεν βρέθηκαν σχετικά traveler data points για το συγκεκριμένο profile._")
        lines.append("")
    else:
        for index, statistic in enumerate(matched_statistics[:12], start=1):
            source = statistic.get("source", {}) or {}
            parent = statistic.get("parent_research", {}) or {}
            lines.append(f"### {index}. {statistic.get('stat', '')}")
            lines.append("")
            if statistic.get("behavior_type"):
                lines.append(f"- **Type:** {statistic.get('behavior_type')}")
            if statistic.get("what_this_shows"):
                lines.append(f"- **What this shows:** {statistic.get('what_this_shows')}")
            if statistic.get("why_it_matches"):
                lines.append(f"- **Why it matches:** {statistic.get('why_it_matches')}")
            if statistic.get("evidence"):
                lines.append(f"- **Evidence:** {statistic.get('evidence')}")
            lines.append(f"- **Source:** {source_label(source)}")
            if parent.get("research_title"):
                lines.append(f"- **Linked research block:** {parent.get('research_title')}")
            lines.append("")

    lines.append("---")
    lines.append("")
    lines.append("## 2. Related Research Context")
    lines.append("")

    if not matches:
        lines.append("_Δεν βρέθηκαν σχετικά research blocks._")
    else:
        for index, match in enumerate(matches[:8], start=1):
            source = match.get("source", {}) or {}
            lines.append(f"### {index}. {match.get('research_title', '')}")
            lines.append("")
            if match.get("why_it_matches"):
                lines.append(f"- **Why it matches:** {match.get('why_it_matches')}")
            if match.get("key_finding"):
                lines.append(f"- **Key finding:** {match.get('key_finding')}")
            if match.get("evidence"):
                lines.append(f"- **Evidence:** {match.get('evidence')}")
            lines.append(f"- **Source:** {source_label(source)}")
            lines.append("")

    return "\n".join(lines)

def _xml_escape(value: Any) -> str:
    """
    Escape text for ReportLab Paragraph XML-like markup.
    """
    text = str(value or "")
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _register_pdf_fonts():
    """
    Register a Unicode font for PDF export.
    This is needed for Greek text. It tries common Windows and Linux font paths.
    """
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    regular_candidates = [
        BASE_DIR / "fonts" / "DejaVuSans.ttf",
        BASE_DIR / "fonts" / "Arial.ttf",
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf"),
    ]
    bold_candidates = [
        BASE_DIR / "fonts" / "DejaVuSans-Bold.ttf",
        BASE_DIR / "fonts" / "Arial Bold.ttf",
        Path("C:/Windows/Fonts/arialbd.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
        Path("/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf"),
    ]

    regular_font = next((path for path in regular_candidates if path.exists()), None)
    bold_font = next((path for path in bold_candidates if path.exists()), None)

    if regular_font:
        try:
            pdfmetrics.registerFont(TTFont("AudienceFont", str(regular_font)))
            if bold_font:
                pdfmetrics.registerFont(TTFont("AudienceFontBold", str(bold_font)))
                return "AudienceFont", "AudienceFontBold"
            return "AudienceFont", "AudienceFont"
        except Exception:
            pass

    # Fallback. This may not render Greek correctly, so installing/reporting a Unicode font is preferred.
    return "Helvetica", "Helvetica-Bold"


def build_audience_intelligence_pdf(
    client_profile: Dict[str, Any],
    matches: List[Dict[str, Any]],
    matched_statistics: List[Dict[str, Any]],
    action_recommendations: Dict[str, List[str]] = None,
) -> bytes:
    """
    Build a simple PDF report focused only on traveler data points and sources.
    No actions, content pillars, digital marketing plan or 30/60/90 plan are included.
    """
    from io import BytesIO
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

    font_regular, font_bold = _register_pdf_fonts()
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=1.6 * cm,
        leftMargin=1.6 * cm,
        topMargin=1.5 * cm,
        bottomMargin=1.5 * cm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleSimple", parent=styles["Title"], fontName=font_bold, fontSize=18, leading=22, alignment=TA_LEFT, textColor=colors.HexColor("#111827"))
    section_style = ParagraphStyle("SectionSimple", parent=styles["Heading2"], fontName=font_bold, fontSize=13, leading=17, spaceBefore=12, spaceAfter=7, textColor=colors.HexColor("#111827"))
    sub_style = ParagraphStyle("SubSimple", parent=styles["Heading3"], fontName=font_bold, fontSize=10.5, leading=14, spaceBefore=8, spaceAfter=4, textColor=colors.HexColor("#991b1b"))
    body_style = ParagraphStyle("BodySimple", parent=styles["BodyText"], fontName=font_regular, fontSize=9, leading=12.5, spaceAfter=5, textColor=colors.HexColor("#1f2937"))
    muted_style = ParagraphStyle("MutedSimple", parent=body_style, fontSize=8.3, textColor=colors.HexColor("#6b7280"))

    def p(value: Any, style=body_style):
        return Paragraph(_xml_escape(value), style)

    def source_label(source: Dict[str, Any]) -> str:
        source = source or {}
        name = str(source.get("source_name", "") or "—").strip()
        report = str(source.get("source_report", "") or "").strip()
        year = str(source.get("source_year", "") or "").strip()
        if report and year:
            return f"{name} - {report} ({year})"
        if report:
            return f"{name} - {report}"
        if year:
            return f"{name} ({year})"
        return name

    def join_values(values: Any) -> str:
        if isinstance(values, list):
            cleaned = [str(value).strip() for value in values if str(value).strip()]
            return ", ".join(cleaned) if cleaned else "Not specified"
        text_value = str(values or "").strip()
        return text_value if text_value else "Not specified"

    client_name = str(client_profile.get("client_name", "Client") or "Client").strip()
    elements = [Paragraph(_xml_escape(f"{client_name} - Traveler Data Points"), title_style)]
    elements.append(p("Research-backed traveler data points and related sources.", muted_style))
    elements.append(Spacer(1, 8))

    profile_rows = [
        [p("Destination", muted_style), p(client_profile.get("main_destination", "Not specified"))],
        [p("Vertical", muted_style), p(client_profile.get("main_vertical", "Not specified"))],
        [p("Products / services", muted_style), p(join_values(client_profile.get("products_services", [])))],
        [p("Target audiences", muted_style), p(join_values(client_profile.get("target_audiences", [])))],
        [p("Markets", muted_style), p(join_values(client_profile.get("markets", [])))],
    ]
    table = Table(profile_rows, colWidths=[4.2 * cm, 11.7 * cm], hAlign="LEFT")
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f8fafc")),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#e5e7eb")),
        ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#e5e7eb")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    elements.append(table)
    elements.append(Spacer(1, 10))

    elements.append(Paragraph("1. Matched Traveler Data Points", section_style))
    if not matched_statistics:
        elements.append(p("No relevant traveler data points were found.", muted_style))
    else:
        for index, statistic in enumerate(matched_statistics[:12], start=1):
            source = statistic.get("source", {}) or {}
            elements.append(Paragraph(_xml_escape(f"{index}. {statistic.get('behavior_type', 'Traveler Behavior')}"), sub_style))
            elements.append(p(f"<b>Data point:</b> {_xml_escape(statistic.get('stat', ''))}"))
            if statistic.get("what_this_shows"):
                elements.append(p(f"<b>What this shows:</b> {_xml_escape(statistic.get('what_this_shows', ''))}"))
            if statistic.get("why_it_matches"):
                elements.append(p(f"<b>Why it matches:</b> {_xml_escape(statistic.get('why_it_matches', ''))}"))
            if statistic.get("evidence"):
                elements.append(p(f"<b>Evidence:</b> {_xml_escape(statistic.get('evidence', ''))}"))
            elements.append(p(f"<b>Source:</b> {_xml_escape(source_label(source))}"))
            elements.append(Spacer(1, 5))

    elements.append(Paragraph("2. Related Research", section_style))
    if not matches:
        elements.append(p("No related research context found.", muted_style))
    else:
        for index, match in enumerate(matches[:8], start=1):
            source = match.get("source", {}) or {}
            elements.append(Paragraph(_xml_escape(f"{index}. {match.get('research_title', '')}"), sub_style))
            if match.get("key_finding"):
                elements.append(p(f"<b>Key finding:</b> {_xml_escape(match.get('key_finding', ''))}"))
            if match.get("why_it_matches"):
                elements.append(p(f"<b>Why it matches:</b> {_xml_escape(match.get('why_it_matches', ''))}"))
            elements.append(p(f"<b>Source:</b> {_xml_escape(source_label(source))}"))
            elements.append(Spacer(1, 5))

    doc.build(elements)
    buffer.seek(0)
    return buffer.getvalue()

def _docx_text(value: Any) -> str:
    """Return safe text for DOCX export."""
    return str(value or "").strip()


def _docx_join_values(values: Any) -> str:
    """Join list values for readable DOCX output."""
    if isinstance(values, list):
        cleaned = [_docx_text(value) for value in values if _docx_text(value)]
        return ", ".join(cleaned) if cleaned else "Not specified"
    text = _docx_text(values)
    return text if text else "Not specified"


def _docx_client_friendly_text(value: Any) -> str:
    """Make generated text sound more natural in the DOCX report.

    The app may generate internal phrasing such as "Ο client".
    In client-facing/editable DOCX exports, we rewrite that to more polished
    business language.
    """
    text = _docx_text(value)
    replacements = {
        "Ο client": "Η επιχείρηση",
        "ο client": "η επιχείρηση",
        "Τον client": "Την επιχείρηση",
        "τον client": "την επιχείρηση",
        "Του client": "Της επιχείρησης",
        "του client": "της επιχείρησης",
        "Στον client": "Στην επιχείρηση",
        "στον client": "στην επιχείρηση",
        "Στο client": "Στην επιχείρηση",
        "στο client": "στην επιχείρηση",
        "client πρέπει": "η επιχείρηση πρέπει",
        "client χρειάζεται": "η επιχείρηση χρειάζεται",
        "client μπορεί": "η επιχείρηση μπορεί",
        "client profile": "business profile",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def _docx_add_bullet(document, text: str, bold_label: str = ""):
    """Add a bullet paragraph, optionally with a bold label at the beginning."""
    paragraph = document.add_paragraph(style="List Bullet")
    if bold_label:
        run = paragraph.add_run(bold_label)
        run.bold = True
        paragraph.add_run(text)
    else:
        paragraph.add_run(text)
    return paragraph


def _docx_add_source_box(document, source: Dict[str, Any], source_file: str = ""):
    """Add source details in a small two-column table."""
    table = document.add_table(rows=0, cols=2)
    table.style = "Light Shading Accent 1"

    rows = [
        ("Source", _docx_text(source.get("source_name", "")) or "—"),
        ("Report", _docx_text(source.get("source_report", "")) or "—"),
        ("Year", _docx_text(source.get("source_year", "")) or "—"),
    ]

    for label, value in rows:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value
        for paragraph in row[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    document.add_paragraph("")


def build_audience_intelligence_docx(
    client_profile: Dict[str, Any],
    matches: List[Dict[str, Any]],
    matched_statistics: List[Dict[str, Any]],
    action_recommendations: Dict[str, List[str]] = None,
    action_cards: Optional[List[Dict[str, Any]]] = None,
    audience_segments: Optional[List[Dict[str, Any]]] = None,
    content_angles: Optional[List[Dict[str, Any]]] = None,
    plan_30_60_90: Optional[Dict[str, List[str]]] = None,
) -> bytes:
    """
    Build an editable DOCX report focused only on traveler data points
    and related research sources. No actions, content pillars, digital
    marketing plan or 30/60/90 plan are included.
    """
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    TOURIX_DARK_RED = RGBColor(127, 29, 29)
    TOURIX_RED = RGBColor(185, 28, 28)
    TOURIX_MUTED = RGBColor(107, 114, 128)
    TOURIX_BODY = RGBColor(31, 41, 55)
    TOURIX_LIGHT_RED_HEX = "FEE2E2"
    TOURIX_SOFT_RED_HEX = "FFF1F2"
    TOURIX_BORDER_RED_HEX = "FCA5A5"

    def set_cell_shading(cell, fill: str):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), fill)

    def set_cell_border(cell, color: str = TOURIX_BORDER_RED_HEX, size: str = "8"):
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = tc_pr.first_child_found_in("w:tcBorders")
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            tc_pr.append(tc_borders)
        for edge in ("top", "left", "bottom", "right"):
            tag = f"w:{edge}"
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), size)
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), color)

    def add_small_note(document, text: str):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = TOURIX_MUTED
        return paragraph

    def add_section_intro(document, text: str):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = TOURIX_MUTED
        return paragraph

    def source_inline(source: Dict[str, Any]) -> str:
        source = source or {}
        source_name = _docx_text(source.get("source_name", "")) or "—"
        source_report = _docx_text(source.get("source_report", ""))
        source_year = _docx_text(source.get("source_year", ""))
        if source_report and source_year:
            return f"{source_name} — {source_report} ({source_year})"
        if source_report:
            return f"{source_name} — {source_report}"
        if source_year:
            return f"{source_name} ({source_year})"
        return source_name

    def add_source_line(document, source: Dict[str, Any]):
        paragraph = document.add_paragraph()
        label = paragraph.add_run("Source: ")
        label.bold = True
        label.font.color.rgb = TOURIX_DARK_RED
        value_run = paragraph.add_run(source_inline(source))
        value_run.font.color.rgb = TOURIX_BODY
        return paragraph

    def add_label_value_table(document, rows: List[tuple[str, str]]):
        table = document.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.style = "Table Grid"
        for label, value in rows:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = value or "Not specified"
            cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_shading(cells[0], TOURIX_LIGHT_RED_HEX)
            set_cell_shading(cells[1], TOURIX_SOFT_RED_HEX)
            set_cell_border(cells[0])
            set_cell_border(cells[1])
            for paragraph in cells[0].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = TOURIX_DARK_RED
        document.add_paragraph("")
        return table

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = TOURIX_BODY
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(22)
    styles["Title"].font.color.rgb = TOURIX_DARK_RED
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = TOURIX_DARK_RED
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = TOURIX_RED

    client_name = _docx_text(client_profile.get("client_name", "Client")) or "Client"

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run(f"{client_name} — Traveler Data Points")
    title_run.font.color.rgb = TOURIX_DARK_RED

    accent = document.add_table(rows=1, cols=1)
    accent.alignment = WD_TABLE_ALIGNMENT.LEFT
    accent_cell = accent.rows[0].cells[0]
    accent_cell.text = ""
    set_cell_shading(accent_cell, "DC2626")
    set_cell_border(accent_cell, "DC2626", "4")
    document.add_paragraph("")

    add_small_note(document, "Editable research output generated from the client/partner profile and matched research database.")
    add_small_note(document, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    document.add_paragraph("")

    document.add_heading("Profile Snapshot", level=1)
    add_label_value_table(document, [
        ("Website", _docx_text(client_profile.get("website", "")) or "Not specified"),
        ("Main destination", _docx_text(client_profile.get("main_destination", "")) or "Not specified"),
        ("Main vertical", _docx_text(client_profile.get("main_vertical", "")) or "Not specified"),
        ("Products / services", _docx_join_values(client_profile.get("products_services", []))),
        ("Target audiences", _docx_join_values(client_profile.get("target_audiences", []))),
        ("Markets", _docx_join_values(client_profile.get("markets", []))),
        ("Research needs", _docx_join_values(client_profile.get("research_needs", []))),
    ])

    document.add_heading("1. Matched Traveler Data Points", level=1)
    add_section_intro(document, "Selected data points are matched with the profile and include the source/research they come from.")

    if not matched_statistics:
        document.add_paragraph("No relevant traveler data points were found for this profile.")
    else:
        for index, statistic in enumerate(matched_statistics[:12], start=1):
            source = statistic.get("source", {}) or {}
            behavior_type = _docx_text(statistic.get("behavior_type", "Traveler Behavior"))
            document.add_heading(f"{index}. {behavior_type}", level=2)
            _docx_add_bullet(document, _docx_text(statistic.get("stat", "")), "Data point: ")
            if statistic.get("what_this_shows"):
                _docx_add_bullet(document, _docx_text(statistic.get("what_this_shows", "")), "What this shows: ")
            if statistic.get("why_it_matches"):
                _docx_add_bullet(document, _docx_text(statistic.get("why_it_matches", "")), "Why it matches: ")
            if statistic.get("evidence"):
                _docx_add_bullet(document, _docx_text(statistic.get("evidence", "")), "Evidence: ")
            add_source_line(document, source)

    document.add_heading("2. Related Research", level=1)
    add_section_intro(document, "Research blocks that support the matched data points and can be used for pre-sales, proposals or partner support.")

    if not matches:
        document.add_paragraph("No related research context was found.")
    else:
        for index, match in enumerate(matches[:8], start=1):
            source = match.get("source", {}) or {}
            document.add_heading(f"{index}. {_docx_text(match.get('research_title', ''))}", level=2)
            if match.get("key_finding"):
                _docx_add_bullet(document, _docx_text(match.get("key_finding", "")), "Key finding: ")
            if match.get("why_it_matches"):
                _docx_add_bullet(document, _docx_text(match.get("why_it_matches", "")), "Why it matches: ")
            if match.get("evidence"):
                _docx_add_bullet(document, _docx_text(match.get("evidence", "")), "Evidence: ")
            add_source_line(document, source)

    document.add_heading("3. Source List", level=1)
    source_rows = []
    seen_sources = set()
    for statistic in matched_statistics[:12]:
        source = statistic.get("source", {}) or {}
        key = source_inline(source)
        if key not in seen_sources:
            seen_sources.add(key)
            source_rows.append(key)
    for match in matches[:8]:
        source = match.get("source", {}) or {}
        key = source_inline(source)
        if key not in seen_sources:
            seen_sources.add(key)
            source_rows.append(key)

    if source_rows:
        table = document.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Source / report"
        set_cell_shading(hdr[0], TOURIX_LIGHT_RED_HEX)
        set_cell_border(hdr[0])
        for paragraph in hdr[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = TOURIX_DARK_RED
        for row_index, source_text in enumerate(source_rows):
            cells = table.add_row().cells
            cells[0].text = source_text
            set_cell_shading(cells[0], TOURIX_SOFT_RED_HEX if row_index % 2 == 0 else "FFFFFF")
            set_cell_border(cells[0])
    else:
        document.add_paragraph("No supporting sources found.")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def render_client_insights_generator():
    st.header("Client Research Data Points")
    st.caption(
        "Simple workspace for matching a client or partner profile with traveler data points and related research sources."
    )

    st.markdown(
        """
        <style>
        .cai-hero {
            border: 1px solid #e5e7eb;
            border-radius: 18px;
            padding: 20px 22px;
            background: linear-gradient(135deg, #ffffff 0%, #f8fafc 100%);
            margin-bottom: 18px;
            box-shadow: 0 1px 2px rgba(16, 24, 40, 0.04);
        }
        .cai-metric {
            border: 1px solid #e5e7eb;
            border-radius: 16px;
            padding: 15px 17px;
            background: #ffffff;
            min-height: 94px;
            box-shadow: 0 1px 3px rgba(16, 24, 40, 0.05);
        }
        .cai-metric-value {
            font-size: 1.8rem;
            font-weight: 850;
            color: #111827;
            line-height: 1.1;
        }
        .cai-metric-label {
            font-size: 0.84rem;
            font-weight: 800;
            color: #111827;
            margin-top: 8px;
        }
        .cai-metric-sub {
            font-size: 0.76rem;
            color: #6b7280;
            margin-top: 4px;
        }
        .cai-card {
            border: 1px solid #e5e7eb;
            border-radius: 18px;
            padding: 18px 20px;
            margin: 10px 0 16px 0;
            background: #ffffff;
            box-shadow: 0 1px 3px rgba(16, 24, 40, 0.04);
        }
        .cai-card-blue {
            border: 1px solid #bfdbfe;
            background: #eff6ff;
        }
        .cai-card-green {
            border: 1px solid #bbf7d0;
            background: #f0fdf4;
        }
        .cai-card-amber {
            border: 1px solid #fde68a;
            background: #fffbeb;
        }
        .cai-card-purple {
            border: 1px solid #ddd6fe;
            background: #faf5ff;
        }
        .cai-mini-title {
            font-size: 0.78rem;
            font-weight: 850;
            text-transform: uppercase;
            letter-spacing: 0.055em;
            color: #475569;
            margin-bottom: 8px;
        }
        .cai-main-text {
            font-size: 1.04rem;
            font-weight: 780;
            color: #111827;
            line-height: 1.48;
            margin-bottom: 8px;
        }
        .cai-muted {
            color: #64748b;
            font-size: 0.86rem;
            line-height: 1.48;
        }
        .cai-source {
            border-left: 4px solid #64748b;
            padding: 10px 12px;
            background: #f8fafc;
            border-radius: 10px;
            margin-top: 10px;
            font-size: 0.88rem;
            color: #334155;
        }
        .cai-pill {
            display: inline-block;
            padding: 5px 10px;
            border-radius: 999px;
            font-weight: 850;
            font-size: 0.72rem;
            margin: 0 6px 8px 0;
            background: #eef2ff;
            color: #3730a3;
        }
        .cai-pill-green { background: #dcfce7; color: #166534; }
        .cai-pill-orange { background: #ffedd5; color: #c2410c; }
        .cai-pill-purple { background: #f3e8ff; color: #7e22ce; }
        .cai-pill-gray { background: #f1f5f9; color: #475569; }
        .cai-step {
            border-left: 4px solid #111827;
            background: #f8fafc;
            padding: 12px 14px;
            border-radius: 12px;
            margin: 8px 0 10px 0;
        }
        .stButton > button {
            border-radius: 12px;
            font-weight: 750;
        }
        div[role="radiogroup"] label {
            border: 1px solid #e5e7eb !important;
            border-radius: 999px !important;
            padding: 7px 12px !important;
            margin-right: 6px !important;
            background: #ffffff !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    def _as_list(value: Any) -> List[str]:
        if value is None:
            return []
        if isinstance(value, list):
            return [str(item).strip() for item in value if str(item).strip()]
        if isinstance(value, dict):
            return [str(item).strip() for item in value.values() if str(item).strip()]
        text = str(value).strip()
        return [text] if text else []

    def _short(text: Any, limit: int = 170) -> str:
        text = str(text or "").strip()
        if len(text) <= limit:
            return text
        return text[: limit - 1].rstrip() + "…"


    def _is_percentage_stat(statistic: Dict[str, Any]) -> bool:
        """Return True only for statistics that contain an actual percentage.

        Used in the Statistics section so that this tab stays focused on
        percentage-based data points, not market-size forecasts or generic findings.
        """
        import re

        fields_to_check = [
            statistic.get("stat", ""),
            statistic.get("evidence", ""),
            statistic.get("what_this_shows", ""),
        ]
        text = " ".join(str(value or "") for value in fields_to_check).lower()

        if "%" in text:
            return True

        # Covers English percentage wording if a source does not use the % symbol.
        if re.search(r"\b\d+(?:[\.,]\d+)?\s*(percent|percentage|pct)\b", text):
            return True

        # Covers Greek percentage wording, e.g. "60 τοις εκατό".
        if re.search(r"\b\d+(?:[\.,]\d+)?\s*(τοις εκατό|ποσοστό|ποσοστ)", text):
            return True

        return False

    def _source_label(source: Dict[str, Any]) -> str:
        source = source or {}
        name = str(source.get("source_name", "") or "—").strip()
        report = str(source.get("source_report", "") or "").strip()
        year = str(source.get("source_year", "") or "").strip()
        if report and year:
            return f"{name} — {report} ({year})"
        if report:
            return f"{name} — {report}"
        if year:
            return f"{name} ({year})"
        return name

    def _source_file(source: Dict[str, Any]) -> str:
        return str((source or {}).get("source_file", "") or "—").strip()

    def _metric_card(value: Any, label: str, sub: str):
        st.markdown(
            f"""
            <div class="cai-metric">
                <div class="cai-metric-value">{value}</div>
                <div class="cai-metric-label">{label}</div>
                <div class="cai-metric-sub">{sub}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    def _source_box(source: Dict[str, Any]):
        st.markdown(
            f"""
            <div class="cai-source">
                <strong>Source</strong><br>
                {_source_label(source)}<br>
                <span class="cai-muted">Source file: {_source_file(source)}</span>
            </div>
            """,
            unsafe_allow_html=True,
        )

    def _research_source_popover(
        title: str,
        source: Dict[str, Any],
        supporting_text: str = "",
        evidence: str = "",
        linked_research: str = "",
        matched_tags: Optional[Dict[str, Any]] = None,
    ):
        """Show the research/source behind a statistic, action or content angle.

        Uses Streamlit popover when available, otherwise falls back to an expander.
        This keeps the UI light: the user can open the research only when needed.
        """
        label = "🔎 Δες την έρευνα / πηγή"
        container = st.popover(label, use_container_width=False) if hasattr(st, "popover") else st.expander(label, expanded=False)

        with container:
            st.markdown(f"#### {title or 'Research source'}")

            if supporting_text:
                st.markdown("**Στοιχείο / finding**")
                st.write(supporting_text)

            if evidence:
                st.markdown("**Evidence**")
                st.write(evidence)

            if linked_research:
                st.markdown("**Linked research block**")
                st.write(linked_research)

            st.markdown("**Source details**")
            _source_box(source or {})

            if matched_tags:
                with st.expander("Matched tags"):
                    st.json(matched_tags)

    def _render_stat_card(statistic: Dict[str, Any], index: int, expanded: bool = False):
        source = statistic.get("source", {}) or {}
        parent = statistic.get("parent_research", {}) or {}
        behavior = statistic.get("behavior_type", "Audience Behavior")
        score = statistic.get("statistic_score", 0)

        if "Planning" in behavior or "Booking" in behavior:
            pill_class = "cai-pill-orange"
        elif "Inspiration" in behavior or "Social" in behavior:
            pill_class = "cai-pill-purple"
        else:
            pill_class = "cai-pill-green"

        st.markdown(
            f"""
            <div class="cai-card">
                <span class="cai-pill {pill_class}">{behavior}</span>
                <span class="cai-pill">{score}/100</span>
                <div class="cai-main-text">{index}. {statistic.get('stat', '')}</div>
                <div class="cai-mini-title">What this means for the client</div>
                <div class="cai-muted">{statistic.get('client_implication', '')}</div>
                <div class="cai-source">
                    <strong>Source</strong><br>
                    {_source_label(source)}<br>
                    <span class="cai-muted">Source file: {_source_file(source)}</span>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
        _research_source_popover(
            title=f"Statistic {index}: source research",
            source=source,
            supporting_text=str(statistic.get("stat", "")),
            evidence=str(statistic.get("evidence", "")),
            linked_research=str(parent.get("research_title", "")),
            matched_tags=statistic.get("matched_tags", {}) or {},
        )

        with st.expander(f"Open evidence, rationale & matched tags — statistic {index}", expanded=expanded):
            col_a, col_b = st.columns(2)
            with col_a:
                st.markdown("**What this shows**")
                st.write(statistic.get("what_this_shows", ""))
                st.markdown("**Why it was selected**")
                st.write(statistic.get("why_it_matches", ""))
            with col_b:
                st.markdown("**Evidence**")
                st.write(statistic.get("evidence", ""))
                st.markdown("**Linked research block**")
                st.write(parent.get("research_title", "—"))
            st.markdown("**Matched tags**")
            st.json(statistic.get("matched_tags", {}) or {})

    uploaded_client_file = st.file_uploader(
        "Upload client profile JSON",
        type=["json"],
        key="client_profile_upload",
        help="Ανέβασε το JSON που παράγεται από το client profile form.",
    )

    if uploaded_client_file is not None:
        upload_signature = f"{uploaded_client_file.name}:{getattr(uploaded_client_file, 'size', 0)}"

        # Streamlit reruns the script every time the user changes section.
        # The uploaded file remains available on rerun, so without this guard the app
        # would reload the JSON and clear the generated insights whenever a section is clicked.
        if st.session_state.get("client_profile_upload_signature") != upload_signature:
            try:
                client_profile = load_client_profile(uploaded_client_file)
                missing_fields = validate_client_profile(client_profile)

                st.session_state["client_profile"] = client_profile
                st.session_state["client_profile_upload_signature"] = upload_signature
                st.session_state["client_insight_matches"] = []
                st.session_state["client_matched_statistics"] = []
                st.session_state["client_action_recommendations"] = {}
                st.session_state["client_action_cards"] = []
                st.session_state["client_audience_segments"] = []
                st.session_state["client_content_angles"] = []
                st.session_state["client_30_60_90_plan"] = {}
                st.session_state["cai_section"] = "🏠 Overview"

                st.success("Client profile loaded successfully.")
                if missing_fields:
                    st.warning("Some important fields are missing or empty: " + ", ".join(missing_fields))

            except ValueError as error:
                st.error(str(error))

    if "client_profile" in st.session_state:
        client_profile = st.session_state["client_profile"]
        client_name = client_profile.get("client_name", "")
        main_vertical = client_profile.get("main_vertical", "")
        main_destination = client_profile.get("main_destination", "")
        products = client_profile.get("products_services", []) or []
        audiences = client_profile.get("target_audiences", []) or []
        goals = client_profile.get("business_goals", []) or []
        channels = client_profile.get("channels", []) or []

        st.markdown("### Client profile snapshot")
        st.markdown(
            f"""
            <div class="cai-hero">
                <span class="cai-pill cai-pill-gray">Client</span>
                <span class="cai-pill cai-pill-gray">{main_vertical or 'No vertical'}</span>
                <span class="cai-pill cai-pill-gray">{main_destination or 'No destination'}</span>
                <div style="font-size:1.55rem;font-weight:850;color:#111827;margin-top:8px;">{client_name or 'Client profile loaded'}</div>
                <div class="cai-muted" style="margin-top:6px;">
                    {_short(', '.join(_as_list(audiences)), 180) or 'No audiences selected'}
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        profile_cols = st.columns(4)
        profile_cols[0].metric("Products / services", len(_as_list(products)))
        profile_cols[1].metric("Audiences", len(_as_list(audiences)))
        profile_cols[2].metric("Goals", len(_as_list(goals)))
        profile_cols[3].metric("Channels", len(_as_list(channels)))

        with st.expander("Open full client profile JSON"):
            st.json(client_profile)

        if st.button("✨ Generate traveler data points", use_container_width=True, type="primary"):
            try:
                research_blocks = load_research_blocks()
                matches = match_research_to_client(client_profile, research_blocks)
                matched_statistics = match_statistics_to_client(client_profile, research_blocks)
                action_recommendations = {}
                action_cards = []
                audience_segments = analyze_best_matching_audience_segments(
                    client_profile,
                    matched_statistics,
                )
                content_angles = []
                plan_30_60_90 = {}

                st.session_state["client_insight_matches"] = matches
                st.session_state["client_matched_statistics"] = matched_statistics
                st.session_state["client_action_recommendations"] = action_recommendations
                st.session_state["client_action_cards"] = action_cards
                st.session_state["client_audience_segments"] = audience_segments
                st.session_state["client_content_angles"] = content_angles
                st.session_state["client_30_60_90_plan"] = plan_30_60_90
                st.session_state["cai_section"] = "🏠 Overview"
                st.session_state["cai_section_radio"] = "🏠 Overview"

            except FileNotFoundError:
                st.error("Missing data/research_blocks.json file.")
            except Exception as error:
                st.error(f"Could not generate audience intelligence: {error}")

    matches = st.session_state.get("client_insight_matches", [])
    matched_statistics = st.session_state.get("client_matched_statistics", [])
    action_recommendations = st.session_state.get("client_action_recommendations", {})
    action_cards = st.session_state.get("client_action_cards", [])
    audience_segments = st.session_state.get("client_audience_segments", [])
    content_angles = st.session_state.get("client_content_angles", [])
    plan_30_60_90 = st.session_state.get("client_30_60_90_plan", {})

    if not (matches or matched_statistics):
        if "client_profile" in st.session_state:
            st.info("Click Generate traveler data points to create the research matcher output.")
        return

    def _build_source_rows() -> List[Dict[str, str]]:
        rows: List[Dict[str, str]] = []
        seen = set()
        for statistic in matched_statistics[:10]:
            source = statistic.get("source", {}) or {}
            key = (_source_label(source), _source_file(source))
            if key not in seen:
                seen.add(key)
                rows.append({"Source / report": key[0], "Source file": key[1]})
        for match in matches[:8]:
            source = match.get("source", {}) or {}
            key = (_source_label(source), _source_file(source))
            if key not in seen:
                seen.add(key)
                rows.append({"Source / report": key[0], "Source file": key[1]})
        return rows

    planning_statistics = [stat for stat in matched_statistics if stat.get("behavior_type") == "Planning & Booking"]
    action_count = sum(len(actions) for actions in action_recommendations.values()) if action_recommendations else 0
    top_stat_score = matched_statistics[0].get("statistic_score", 0) if matched_statistics else 0
    strongest_segment = audience_segments[0].get("audience_segment", "—") if audience_segments else "—"

    st.divider()
    st.markdown("## Research data points overview")
    source_count = len(_build_source_rows())
    metric_cols = st.columns(3)
    with metric_cols[0]:
        _metric_card(len(matched_statistics), "Traveler data points", "Matched to this client / partner profile")
    with metric_cols[1]:
        _metric_card(len(planning_statistics), "Planning / booking signals", "Decision journey evidence")
    with metric_cols[2]:
        _metric_card(source_count, "Related research sources", "Reports used in this output")

    section_options = [
        "🏠 Overview",
        "📊 Statistics",
        "🧭 Journey",
        "👥 Audiences",
        "📚 Sources",
        "📤 Export",
    ]
    default_section = st.session_state.get("cai_section", "🏠 Overview")
    if default_section not in section_options:
        default_section = "🏠 Overview"

    section = st.radio(
        "Choose what you want to explore",
        section_options,
        index=section_options.index(default_section),
        horizontal=True,
        label_visibility="collapsed",
        key="cai_section_radio",
    )
    st.session_state["cai_section"] = section

    if section == "🏠 Overview":
        st.markdown("### Executive snapshot")
        left, right = st.columns([1.35, 1])
        with left:
            st.markdown("#### Top audience behavior takeaways")
            for index, statistic in enumerate(matched_statistics[:3], start=1):
                _render_stat_card(statistic, index, expanded=False)
        with right:
            st.markdown("#### Quick research focus")
            if matched_statistics:
                top_source = matched_statistics[0].get("source", {}) or {}
                st.markdown(
                    f"""
                    <div class="cai-card cai-card-blue">
                        <div class="cai-mini-title">Strongest matched source</div>
                        <div class="cai-main-text">{_source_label(top_source)}</div>
                        <div class="cai-muted">Use this as the first research reference for pre-sales, proposal or partner support.</div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
            st.markdown(
                f"""
                <div class="cai-card cai-card-green">
                    <div class="cai-mini-title">Strongest traveler segment</div>
                    <div class="cai-main-text">{strongest_segment}</div>
                    <div class="cai-muted">Based on the matched traveler data points and research tags.</div>
                </div>
                """,
                unsafe_allow_html=True,
            )

    elif section == "📊 Statistics":
        st.markdown("### Key audience behavior statistics")
        st.caption(
            "Εδώ εμφανίζονται μόνο statistics που περιέχουν ποσοστό (%), ώστε το section να μένει καθαρό και πραγματικά data-driven."
        )

        percentage_statistics = [
            statistic for statistic in matched_statistics
            if _is_percentage_stat(statistic)
        ]

        if not percentage_statistics:
            st.info(
                "Δεν βρέθηκαν statistics με ποσοστά για το συγκεκριμένο client profile. "
                "Τα υπόλοιπα findings παραμένουν διαθέσιμα στα Journey και Actions sections."
            )
        else:
            labels = [
                f"{idx}. {stat.get('behavior_type', 'Audience')} · {stat.get('statistic_score', 0)}/100 · {_short(stat.get('stat', ''), 70)}"
                for idx, stat in enumerate(percentage_statistics[:10], start=1)
            ]
            selected_label = st.selectbox("Choose percentage statistic", labels, label_visibility="collapsed")
            selected_index = labels.index(selected_label)
            selected_stat = percentage_statistics[selected_index]
            _render_stat_card(selected_stat, selected_index + 1, expanded=True)

            st.markdown("#### Quick scan — percentage statistics only")
            for idx, statistic in enumerate(percentage_statistics[:10], start=1):
                source = statistic.get("source", {}) or {}
                with st.expander(
                    f"{idx}. {statistic.get('behavior_type', 'Audience')} · {statistic.get('statistic_score', 0)}/100 — {_short(statistic.get('stat', ''), 95)}"
                ):
                    st.markdown("**Statistic**")
                    st.write(statistic.get("stat", ""))
                    st.markdown("**Client implication**")
                    st.write(statistic.get("client_implication", ""))
                    _research_source_popover(
                        title=f"Percentage statistic {idx}: source research",
                        source=source,
                        supporting_text=str(statistic.get("stat", "")),
                        evidence=str(statistic.get("evidence", "")),
                        linked_research=str((statistic.get("parent_research", {}) or {}).get("research_title", "")),
                        matched_tags=statistic.get("matched_tags", {}) or {},
                    )
                    _source_box(source)

    elif section == "🧭 Journey":
        st.markdown("### Journey-related traveler signals")
        journey_stage = st.radio(
            "Journey stage",
            ["Inspiration", "Planning", "Booking", "Post-experience"],
            horizontal=True,
            label_visibility="collapsed",
        )

        if journey_stage == "Inspiration":
            relevant_stats = [s for s in matched_statistics if "Inspiration" in s.get("behavior_type", "") or "Social" in s.get("behavior_type", "")]
        elif journey_stage == "Planning":
            relevant_stats = [s for s in matched_statistics if "Planning" in s.get("behavior_type", "")]
        elif journey_stage == "Booking":
            relevant_stats = [s for s in matched_statistics if "Booking" in s.get("behavior_type", "")]
        else:
            relevant_stats = [s for s in matched_statistics if "Trust" in s.get("behavior_type", "") or "Value" in s.get("behavior_type", "")]

        st.markdown(f"#### Relevant traveler data points for {journey_stage}")
        if relevant_stats:
            for idx, statistic in enumerate(relevant_stats[:6], start=1):
                _render_stat_card(statistic, idx, expanded=False)
        else:
            st.info("No specific traveler data points found for this stage.")

    elif section == "👥 Audiences":
        st.markdown("### Audience explorer")
        if audience_segments:
            segment_labels = [segment.get("audience_segment", "Segment") for segment in audience_segments[:10]]
            selected_segment_label = st.radio("Audience segment", segment_labels, horizontal=True, label_visibility="collapsed")
            segment = next((s for s in audience_segments if s.get("audience_segment") == selected_segment_label), audience_segments[0])
            st.markdown(
                f"""
                <div class="cai-card cai-card-purple">
                    <span class="cai-pill cai-pill-purple">High priority</span>
                    <div class="cai-main-text">{segment.get('audience_segment', '')}</div>
                    <div class="cai-muted">{segment.get('matched_count', segment.get('score', '—'))} matched statistics · Avg score {segment.get('average_score', segment.get('avg_score', '—'))}/100</div>
                    <div class="cai-mini-title" style="margin-top:12px;">Why it matters</div>
                    <div class="cai-muted">{segment.get('why_it_matters', segment.get('reason', '')) or 'This segment appears in the matched research database for this client profile.'}</div>
                </div>
                """,
                unsafe_allow_html=True,
            )
            with st.expander("Open segment details and source references"):
                sources = segment.get("sources", []) or []
                if sources:
                    st.markdown("**Research sources supporting this audience segment**")
                    for src in sources:
                        st.markdown(f"- {src}")
                st.json(segment)
        else:
            st.info("No audience segments were generated.")

    elif section == "📚 Sources":
        st.markdown("### Related research sources")
        source_rows = _build_source_rows()
        if source_rows:
            st.dataframe(pd.DataFrame(source_rows), width="stretch", hide_index=True)
        else:
            st.info("No related sources found.")

        st.markdown("### Related research context")
        if matches:
            for index, match in enumerate(matches[:8], start=1):
                source = match.get("source", {}) or {}
                with st.expander(f"{index}. {match.get('research_title', '')}"):
                    st.markdown("**Why it matches**")
                    st.write(match.get("why_it_matches", ""))
                    st.markdown("**Key finding**")
                    st.write(match.get("key_finding", ""))
                    st.markdown("**Evidence**")
                    st.write(match.get("evidence", ""))
                    _source_box(source)
        else:
            st.info("No related research context found.")

    elif section == "📤 Export":
        st.markdown("### Export reports and data")
        client_profile = st.session_state.get("client_profile", {})
        safe_client_name = (
            str(client_profile.get("client_name", "client"))
            .strip()
            .replace(" ", "_")
            .replace("/", "_")
            .replace("\\", "_")
        ) or "client"

        markdown_report = build_audience_intelligence_markdown(
            client_profile=client_profile,
            matches=matches,
            matched_statistics=matched_statistics,
            action_recommendations=action_recommendations,
            action_cards=action_cards,
            audience_segments=audience_segments,
            content_angles=content_angles,
            day_plan=plan_30_60_90,
        )
        st.download_button(
            "Download Markdown report",
            data=markdown_report.encode("utf-8"),
            file_name=f"{safe_client_name}_audience_intelligence.md",
            mime="text/markdown",
            use_container_width=True,
        )

        try:
            pdf_report = build_audience_intelligence_pdf(
                client_profile=client_profile,
                matches=matches,
                matched_statistics=matched_statistics,
                action_recommendations=action_recommendations,
            )
            st.download_button(
                "Download PDF report",
                data=pdf_report,
                file_name=f"{safe_client_name}_audience_intelligence.pdf",
                mime="application/pdf",
                use_container_width=True,
            )
        except ModuleNotFoundError:
            st.warning("PDF export requires reportlab. Run: pip install reportlab")
        except Exception as pdf_error:
            st.warning(f"Could not generate PDF report: {pdf_error}")

        try:
            docx_report = build_audience_intelligence_docx(
                client_profile=client_profile,
                matches=matches,
                matched_statistics=matched_statistics,
                action_recommendations=action_recommendations,
                action_cards=action_cards,
                audience_segments=audience_segments,
                content_angles=content_angles,
                plan_30_60_90=plan_30_60_90,
            )
            st.download_button(
                "Download editable DOCX report",
                data=docx_report,
                file_name=f"{safe_client_name}_audience_intelligence.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except ModuleNotFoundError:
            st.warning("DOCX export requires python-docx. Run: pip install python-docx")
        except Exception as docx_error:
            st.warning(f"Could not generate DOCX report: {docx_error}")

        enhanced_output = {
            "client_profile": client_profile,
            "traveler_data_points": matched_statistics[:12],
            "related_research_context": matches[:10],
            "audience_segments": audience_segments,
        }
        st.download_button(
            "Download research data points JSON",
            data=json.dumps(enhanced_output, ensure_ascii=False, indent=2).encode("utf-8"),
            file_name="client_research_data_points.json",
            mime="application/json",
            use_container_width=True,
        )
        if matched_statistics:
            st.download_button(
                "Download audience statistics JSON",
                data=json.dumps(matched_statistics[:10], ensure_ascii=False, indent=2).encode("utf-8"),
                file_name="client_audience_statistics.json",
                mime="application/json",
                use_container_width=True,
            )
        if matches:
            st.download_button(
                "Download supporting context JSON",
                data=json.dumps(matches[:10], ensure_ascii=False, indent=2).encode("utf-8"),
                file_name="client_supporting_research_context.json",
                mime="application/json",
                use_container_width=True,
            )

def main():
    st.set_page_config(page_title=APP_TITLE, layout="wide")

    if not require_password():
        return

    st.title(APP_TITLE)
    st.caption(APP_SUBTITLE)

    extractor_tab, insights_tab = st.tabs([
        "Research Extractor",
        "Client Research Matcher",
    ])

    with extractor_tab:
        render_research_extractor()

    with insights_tab:
        render_client_insights_generator()


def render_research_extractor():
    for key, default in {
        "last_extraction_result": None,
        "last_effective_source_url": "",
        "last_loaded_label": "",
        "last_source_type": "",
        "last_pdf_source_url": "",
        "last_pdf_file_name": "",
        "last_sent_teable_record_id": "",
        "send_anyway_duplicate": False,
    }.items():
        if key not in st.session_state:
            st.session_state[key] = default

    gemini_key = get_secret("GEMINI_API_KEY")
    gemini_model = get_secret("GEMINI_MODEL", "gemini-2.5-flash-lite")
    gemini_fallback_model = get_secret("GEMINI_FALLBACK_MODEL", "gemini-2.5-flash")
    teable_token = get_secret("TEABLE_API_TOKEN")
    teable_table_id = get_secret("TEABLE_TABLE_ID")

    with st.sidebar:
        st.header("Input")
        uploaded_pdf = st.file_uploader("Upload PDF", type=["pdf"])
        source_url = st.text_input("Or paste a URL")
        pasted_article_text = st.text_area(
            "Or paste article text",
            height=180,
            placeholder="Use this when a website blocks automated access or copy/paste is easier.",
        )
        run = st.button("Run extraction", use_container_width=True)

        st.divider()
        if gemini_key and teable_token and teable_table_id:
            st.success("System ready")
        elif gemini_key:
            st.warning("Gemini ready, Teable missing")
        else:
            st.warning("Gemini key missing")

        with st.expander("What if URL is blocked?"):
            st.write(
                "If a website blocks automated access, save the page as PDF and upload it, "
                "or paste the article text manually. Image-based PDFs will use OCR/vision fallback."
            )

    if run:
        if not uploaded_pdf and not source_url and not pasted_article_text.strip():
            st.warning("Upload a PDF, paste article text, or enter a URL, then click Run extraction.")
            st.stop()

        pages = []
        loaded_label = ""
        effective_source_url = source_url.strip()
        source_type = ""
        pdf_source_url = ""
        pdf_file_name = ""

        try:
            if uploaded_pdf is not None:
                with st.spinner("Reading PDF..."):
                    pages = extract_pdf_pages(uploaded_pdf)
                source_type = "PDF upload"
                pdf_file_name = uploaded_pdf.name
                pdf_source_url = source_url.strip()
                loaded_label = f"Loaded: {uploaded_pdf.name} ({len(pages)} pages)"
            elif pasted_article_text.strip():
                pages = extract_pasted_article_pages(pasted_article_text)
                effective_source_url = source_url.strip()
                source_type = "Pasted article text"
                pdf_source_url = ""
                pdf_file_name = ""
                loaded_label = "Loaded: pasted article text"
            else:
                with st.spinner("Fetching URL..."):
                    fetched = fetch_url(effective_source_url)

                effective_source_url = fetched.get("source_url", effective_source_url)

                if fetched["type"] == "pdf":
                    pages = extract_pdf_pages(fetched["content"])
                    source_type = "URL PDF"
                    pdf_source_url = effective_source_url
                    pdf_file_name = ""
                    loaded_label = f"Loaded: {effective_source_url} ({len(pages)} pages)"
                else:
                    pages = extract_html_pages(fetched["content"])
                    source_type = "URL web page"
                    pdf_source_url = ""
                    pdf_file_name = ""
                    loaded_label = f"Loaded: {effective_source_url} (web page)"

            st.info(loaded_label)

        except FetchBlockedError as e:
            st.error(str(e))
            st.info(
                "This website blocks automated access. Open the article in your browser, "
                "copy the article text into 'Or paste article text', or save the page as PDF and upload it."
            )
            st.stop()
        except FetchParseError as e:
            st.error(str(e))
            st.stop()
        except Exception as e:
            st.error(f"Input processing failed: {e}")
            st.stop()

        try:
            with st.spinner("Extracting fields..."):
                result = build_output(
                    pages=pages,
                    source_url=effective_source_url,
                    pdf_path=uploaded_pdf.name if uploaded_pdf is not None else "",
                )
        except Exception as e:
            st.error(f"Extraction failed: {e}")
            st.stop()

        st.session_state.last_extraction_result = result
        st.session_state.last_effective_source_url = effective_source_url
        st.session_state.last_loaded_label = loaded_label
        st.session_state.last_source_type = source_type
        st.session_state.last_pdf_source_url = pdf_source_url
        st.session_state.last_pdf_file_name = pdf_file_name
        st.session_state.last_sent_teable_record_id = ""
        st.session_state.send_anyway_duplicate = False

    result = st.session_state.last_extraction_result

    if result is None:
        st.info("Upload a PDF, paste article text, or paste a URL, then click Run extraction.")
        return

    effective_source_url = st.session_state.last_effective_source_url
    loaded_label = st.session_state.last_loaded_label
    source_type = st.session_state.last_source_type
    pdf_source_url = st.session_state.last_pdf_source_url
    pdf_file_name = st.session_state.last_pdf_file_name
    row = result.final_row
    review_meta = result.review_meta

    if loaded_label:
        st.info(loaded_label)

    if result.llm_error:
        st.error(friendly_llm_error(result.llm_error))

    filled = sum(1 for v in row.values() if str(v).strip().lower() != "not specified")
    total = len(DISPLAY_COLUMNS)
    confs = [meta.get("confidence_pct", 0) for meta in review_meta.values()]
    avg_conf = int(sum(confs) / len(confs)) if confs else 0
    needs_review = sum(1 for meta in review_meta.values() if meta.get("needs_review", True))
    ocr_used = bool((result.llm_debug or {}).get("ocr_used", False))

    status_default = compute_default_status(needs_review)
    quality_default = compute_source_quality(row, avg_conf)

    metadata = {
        "PDF Source URL": pdf_source_url,
        "PDF File Name": pdf_file_name,
        "Source Type": source_type,
        "OCR Used": ocr_used,
        "Gemini Used": bool(result.llm_used),
        "Confidence Score": avg_conf,
        "Needs Review Count": needs_review,
        "Status": status_default,
        "Source Quality": quality_default,
        "Usefulness": generate_default_usefulness(row, {
            "Source Type": source_type,
            "Confidence Score": avg_conf,
            "Needs Review Count": needs_review,
        }),
        "Reviewer Notes": "",
    }

    m1, m2, m3, m4 = st.columns(4)
    with m1:
        st.metric("Fields extracted", f"{filled} / {total}")
    with m2:
        st.metric("Avg. confidence", f"{avg_conf}%")
    with m3:
        st.metric("Needs review", str(needs_review))
    with m4:
        st.metric("OCR used", "Yes" if ocr_used else "No")

    tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs(
        [
            "Review & send",
            "Field-by-field",
            "Spreadsheet row",
            "Field review",
            "Raw preview",
            "Debug",
        ]
    )

    with tab1:
        edited_row = render_editable_review(row, metadata)

        duplicate = None
        if teable_token and teable_table_id:
            with st.spinner("Checking for duplicates in Teable..."):
                duplicate = find_duplicate_in_teable(edited_row, effective_source_url)

        if duplicate:
            st.warning("Possible duplicate found in Teable. Same title or source URL already exists.")
            with st.expander("Duplicate record details"):
                st.json(duplicate)
            st.session_state.send_anyway_duplicate = st.checkbox(
                "Send anyway",
                value=st.session_state.send_anyway_duplicate,
            )

        st.divider()
        send_disabled = bool(duplicate and not st.session_state.send_anyway_duplicate)
        send_to_teable = st.button(
            "Send reviewed result to Teable",
            use_container_width=True,
            disabled=send_disabled,
        )

        if st.session_state.last_sent_teable_record_id:
            st.success(f"Already sent to Teable. Record ID: {st.session_state.last_sent_teable_record_id}")
        elif not teable_token or not teable_table_id:
            st.warning("Teable token or table ID is missing in Streamlit Secrets.")

        if send_to_teable:
            try:
                teable_response = append_result_to_teable(edited_row, effective_source_url)

                record_id = ""
                records = teable_response.get("records") if isinstance(teable_response, dict) else None
                if isinstance(records, list) and records:
                    record_id = str(records[0].get("id", ""))

                st.session_state.last_sent_teable_record_id = record_id or "sent"
                st.success(f"Result sent to Teable{f' — Record ID: {record_id}' if record_id else ''}.")
                with st.expander("Teable API response"):
                    st.json(teable_response)

            except Exception as e:
                st.error(f"Could not send to Teable: {e}")

    with tab2:
        df_field = review_to_dataframe(row)
        st.dataframe(style_not_specified(df_field), width="stretch", hide_index=True)

    with tab3:
        df_row = row_to_dataframe(row)
        st.dataframe(df_row, width="stretch", hide_index=True)

        csv_bytes = df_row.to_csv(index=False).encode("utf-8")
        st.download_button(
            "Download CSV row",
            data=csv_bytes,
            file_name="research_extraction_row.csv",
            mime="text/csv",
        )

        json_bytes = pd.Series(row).to_json(force_ascii=False, indent=2).encode("utf-8")
        st.download_button(
            "Download JSON row",
            data=json_bytes,
            file_name="research_extraction_row.json",
            mime="application/json",
        )

    with tab4:
        df_review = review_meta_to_dataframe(review_meta)
        st.dataframe(df_review, width="stretch", hide_index=True)

    with tab5:
        st.text_area(
            "Raw text preview",
            value=result.raw_text_preview or "",
            height=500,
        )

    with tab6:
        debug_payload = {
            "gemini_key_loaded": bool(gemini_key),
            "gemini_model": gemini_model,
            "gemini_fallback_model": gemini_fallback_model,
            "teable_token_loaded": bool(teable_token),
            "teable_table_id_loaded": bool(teable_table_id),
            "source_type": source_type,
            "pdf_source_url": pdf_source_url,
            "pdf_file_name": pdf_file_name,
            "ocr_used": ocr_used,
            "llm_used": result.llm_used,
            "llm_model_used": result.llm_model_used,
            "llm_error_raw": result.llm_error,
            "llm_error_friendly": friendly_llm_error(result.llm_error) if result.llm_error else "",
            "llm_debug": result.llm_debug,
        }
        st.json(debug_payload)


if __name__ == "__main__":
    main()